"""
Builds the formal capstone document (Capstone_Report.docx) from project artifacts.
Uses python-docx. Pulls real numbers from the saved CSVs and figures.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import csv
import os

BASE = r'D:\Desktop\Capstone'

# -----------------------------------------------------------
# Helpers
# -----------------------------------------------------------
def set_cell_shading(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, size=11, bold=False, italic=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def add_figure(doc, path, caption, width=6.0):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(10)
        cap.paragraph_format.space_after = Pt(12)


# -----------------------------------------------------------
# Read tuning data so the report uses real numbers
# -----------------------------------------------------------
with open(os.path.join(BASE, 'k_tuning_results.csv')) as f:
    rows = list(csv.DictReader(f))

# k=7 was the chosen model
k7 = [r for r in rows if r['k'] == '7'][0]
acc = float(k7['accuracy']); prec = float(k7['precision'])
rec = float(k7['recall']); f1 = float(k7['f1'])

# Confusion matrix counts (verified numerically: TP=136, FN=245, FP=59, TN=2025)
TP, FN, FP, TN = 136, 245, 59, 2025

# -----------------------------------------------------------
# Build the document
# -----------------------------------------------------------
doc = Document()

# ============ TITLE PAGE ============
for _ in range(4):
    doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Predicting Online Shoppers’ Purchase Intention')
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('A k-Nearest Neighbors Classifier Implemented from Scratch')
r.italic = True; r.font.size = Pt(16)

doc.add_paragraph()
doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Capstone Project Report')
r.bold = True; r.font.size = Pt(14)

for _ in range(8):
    doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Submitted: May 2026'); r.font.size = Pt(12)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Dataset: Online Shoppers Purchasing Intention (UCI ML Repository)'); r.font.size = Pt(11); r.italic = True

doc.add_page_break()

# ============ ABSTRACT ============
add_heading(doc, 'Abstract', level=1)
add_para(doc,
    'This capstone investigates whether a k-Nearest Neighbors (kNN) classifier, implemented from '
    'first principles in NumPy, can predict online purchase intent from session-level web analytics. '
    'Using the UCI Online Shoppers Purchasing Intention dataset (12,330 sessions, 17 features, '
    '15.47% positive class), I built a complete pipeline covering exploratory analysis, one-hot '
    'encoding, stratified train/test splitting, train-only standardization, hand-coded kNN with '
    'majority-vote prediction, and exhaustive tuning of k from 1 to 30. The final model (k = 7, '
    'Euclidean distance) achieved 87.67% accuracy, 0.6974 precision, 0.3570 recall, and 0.4722 F1 '
    'on a held-out test set, validated by 5-fold stratified cross-validation. The model '
    'outperformed both a majority-class baseline (zero recall) and a single-feature rule on '
    'PageValues. The dominant predictors were PageValues, BounceRates, visit Month (especially '
    'November), and VisitorType. The most honest finding of the project is that the model’s '
    'low recall (it misses ~64% of actual buyers) is not a coding bug but a structural consequence '
    'of severe class imbalance combined with kNN’s majority-vote behavior — a result this '
    'report deliberately surfaces rather than hides.',
    align='justify')

doc.add_page_break()

# ============ 1. INTRODUCTION ============
add_heading(doc, '1. Introduction', level=1)

add_heading(doc, '1.1 Background', level=2)
add_para(doc,
    'E-commerce platforms generate millions of clickstream sessions per day, yet only a small '
    'fraction culminate in a purchase. Understanding which sessions are likely to convert allows '
    'merchants to allocate marketing spend, personalize on-site interventions (pop-ups, discounts, '
    'live-chat invitations), and forecast revenue. Machine learning offers a tractable way to model '
    'this behavior from observable session attributes such as page visit counts, durations, bounce '
    'rates, and Google Analytics-derived metrics.', align='justify')

add_heading(doc, '1.2 Problem Statement', level=2)
add_para(doc,
    'Given the behavioral and metadata features of a single browsing session, can a classification '
    'model predict whether the session will end in a completed purchase? The task is binary '
    'classification on imbalanced data (15.47% positives), making evaluation more nuanced than a '
    'simple accuracy report.', align='justify')

add_heading(doc, '1.3 Objectives', level=2)
add_bullet(doc, 'Implement the full machine-learning pipeline from scratch in NumPy without using scikit-learn’s estimators, in order to demonstrate genuine algorithmic understanding.')
add_bullet(doc, 'Apply k-Nearest Neighbors to the UCI Online Shoppers dataset and report performance on multiple, complementary metrics (accuracy, precision, recall, F1, specificity).')
add_bullet(doc, 'Identify the optimal value of k through systematic hyperparameter sweep.')
add_bullet(doc, 'Validate the result using 5-fold stratified cross-validation and compare against principled baselines.')
add_bullet(doc, 'Critically discuss the model’s limitations, especially under class imbalance, rather than claim a clean win.')

add_heading(doc, '1.4 Significance', level=2)
add_para(doc,
    'The project demonstrates that a transparent, mathematically simple algorithm can produce '
    'commercially actionable predictions when paired with disciplined preprocessing and honest '
    'evaluation. It also exposes the trade-offs that practitioners face on imbalanced data — '
    'a recurring theme in fraud detection, churn prediction, medical screening, and marketing.',
    align='justify')

doc.add_page_break()

# ============ 2. LITERATURE / BACKGROUND ============
add_heading(doc, '2. Background and Related Work', level=1)

add_heading(doc, '2.1 The k-Nearest Neighbors Algorithm', level=2)
add_para(doc,
    'k-Nearest Neighbors is a non-parametric, instance-based ("lazy") learning algorithm. Given a '
    'query point, it computes the distance to every point in the training set, retains the k closest '
    'neighbors, and assigns the majority class label among them. There is no explicit training step '
    '— the entire training set serves as the model. Strengths: simplicity, no assumptions about '
    'the underlying distribution, and naturally non-linear decision boundaries. Weaknesses: '
    'O(n · d) prediction cost per query, sensitivity to feature scaling, the curse of '
    'dimensionality, and sensitivity to class imbalance.', align='justify')

add_heading(doc, '2.2 The Distance Metric', level=2)
add_para(doc,
    'This project uses Euclidean distance, defined for two d-dimensional points x and y as '
    '√Σ(xᵢ − yᵢ)². Because distance treats every feature equally, '
    'features measured in different units (e.g., page counts vs. bounce-rate proportions) must be '
    'standardized to comparable scale; otherwise, the largest-magnitude feature dominates the '
    'distance and effectively becomes the only feature used.', align='justify')

add_heading(doc, '2.3 Choice of k and the Bias–Variance Trade-off', level=2)
add_para(doc,
    'A small k (e.g., 1) yields a highly flexible decision boundary that follows training noise '
    '(low bias, high variance — overfitting). A large k smooths the boundary, lowering variance '
    'but raising bias toward the majority class. The optimal k is dataset-dependent and must be '
    'searched empirically.', align='justify')

add_heading(doc, '2.4 The Dataset in the Literature', level=2)
add_para(doc,
    'The Online Shoppers Purchasing Intention dataset was introduced by Sakar et al. (2019) for '
    'real-time purchase prediction. Their original work compared multilayer perceptrons and LSTM '
    'recurrent networks. Subsequent benchmark studies have applied logistic regression, decision '
    'trees, random forests, gradient boosting, and SVMs, with reported F1 scores typically in the '
    '0.55–0.65 range when class-imbalance handling is applied. This capstone deliberately uses '
    'a simpler model to make the algorithmic mechanics visible, accepting that the absolute score '
    'will trail tuned ensemble methods.', align='justify')

doc.add_page_break()

# ============ 3. DATASET ============
add_heading(doc, '3. Dataset Description', level=1)

add_heading(doc, '3.1 Source and Composition', level=2)
add_bullet(doc, 'Source: UCI Machine Learning Repository / Kaggle (Henry Sue, 2020 mirror).')
add_bullet(doc, 'Total sessions: 12,330 collected over a one-year period.')
add_bullet(doc, 'Features: 10 numerical and 8 categorical attributes (17 predictors after removing the target).')
add_bullet(doc, 'Target: Revenue (Boolean) — True if the session ended with a purchase.')
add_bullet(doc, 'Class distribution: 10,422 No-Purchase (84.53%) vs. 1,908 Purchase (15.47%) — a 5.46 : 1 imbalance ratio.')

add_heading(doc, '3.2 Feature Categories', level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Category'; hdr[1].text = 'Features'; hdr[2].text = 'Description'
for c in hdr:
    set_cell_shading(c, '1F3A5F')
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

rows_data = [
    ('Page visits', 'Administrative, Informational, ProductRelated', 'Number of pages of each type visited.'),
    ('Time spent', '*_Duration columns', 'Total seconds on each page type.'),
    ('Google Analytics', 'BounceRates, ExitRates, PageValues', 'Site analytics; PageValues correlates with revenue.'),
    ('Calendar / closeness', 'SpecialDay, Month', 'Proximity to holidays; visit month.'),
    ('Technical metadata', 'OperatingSystems, Browser, Region, TrafficType', 'User connection profile.'),
    ('Visitor info', 'VisitorType, Weekend', 'Returning vs. new visitor; weekend flag.'),
]
for cat, feats, desc in rows_data:
    row = table.add_row().cells
    row[0].text = cat; row[1].text = feats; row[2].text = desc

doc.add_paragraph()

add_heading(doc, '3.3 Data Quality', level=2)
add_para(doc,
    'No missing values were detected after loading. All numerical features are non-negative. The '
    'two categorical text columns (Month, VisitorType) were one-hot encoded into 13 binary '
    'indicators, raising the final feature count to 28. The Weekend and Revenue Booleans were cast '
    'to 0 / 1 integers.', align='justify')

doc.add_page_break()

# ============ 4. EXPLORATORY DATA ANALYSIS ============
add_heading(doc, '4. Exploratory Data Analysis', level=1)
add_para(doc,
    'EDA was performed before any modeling to understand class separability, feature relationships, '
    'and temporal patterns. Six dedicated visualizations were produced.', align='justify')

add_heading(doc, '4.1 Class Imbalance', level=2)
add_para(doc,
    'For every one purchaser there are roughly 5.5 non-purchasers. A model that always predicts '
    '“No Purchase” would achieve 84.53% accuracy without learning anything — the '
    'first warning sign that accuracy alone cannot validate this work.', align='justify')
add_figure(doc, os.path.join(BASE, 'fig1_class_distribution.png'),
           'Figure 1. Class distribution showing the 84.5% / 15.5% imbalance.')

add_heading(doc, '4.2 Feature Correlations', level=2)
add_para(doc,
    'BounceRates and ExitRates are correlated at r = 0.913 — effectively duplicates. '
    'ProductRelated and ProductRelated_Duration correlate at r = 0.861. These near-duplicates will '
    'give those behaviors double weight inside the kNN distance calculation, an issue the '
    'Limitations section returns to.', align='justify')
add_figure(doc, os.path.join(BASE, 'fig2_correlation_heatmap.png'),
           'Figure 2. Correlation heatmap of numeric features.')

add_heading(doc, '4.3 PageValues — The Strongest Signal', level=2)
add_para(doc,
    'PageValues is the single most discriminative feature in the dataset. The mean for buyers '
    '(27.26) is 13.8× the mean for non-buyers (1.98); the median for non-buyers is exactly 0. '
    'This kind of class separability is precisely what kNN exploits.', align='justify')
add_figure(doc, os.path.join(BASE, 'fig3_pagevalues_boxplot.png'),
           'Figure 3. PageValues by purchase outcome (full and zoomed views).')

add_heading(doc, '4.4 BounceRates — Engagement vs. Drive-by Traffic', level=2)
add_para(doc,
    'Non-buyers had a mean BounceRate roughly five times that of buyers (0.0253 vs. 0.0051). The '
    'median BounceRate for buyers is zero — buyers stay and engage; non-buyers drift in and out.',
    align='justify')
add_figure(doc, os.path.join(BASE, 'fig4_bouncerates_boxplot.png'),
           'Figure 4. BounceRates by purchase outcome.')

add_heading(doc, '4.5 Temporal Effects', level=2)
add_para(doc,
    'Conversion rate varies dramatically across months: from 1.63% in February to 25.35% in '
    'November. The November peak is consistent with end-of-year holiday shopping. The model can '
    'leverage this through the one-hot Month features.', align='justify')
add_figure(doc, os.path.join(BASE, 'fig5_purchases_by_month.png'),
           'Figure 5. Sessions and conversion rate by month.')

add_heading(doc, '4.6 Visitor Type — A Counter-intuitive Result', level=2)
add_para(doc,
    'New visitors convert at 24.91%, while returning visitors convert at only 13.93%. New visitors '
    'buy at almost twice the rate of returning visitors. This is contrary to the intuition that '
    'returning customers convert better and is one of the most interesting findings of the EDA.',
    align='justify')
add_figure(doc, os.path.join(BASE, 'fig6_visitor_type.png'),
           'Figure 6. Sessions and conversion rate by visitor type.')

doc.add_page_break()

# ============ 5. METHODOLOGY ============
add_heading(doc, '5. Methodology', level=1)

add_heading(doc, '5.1 Overall Pipeline', level=2)
add_para(doc, 'The project follows a standard supervised-learning pipeline, executed in seven discrete scripts:')
add_bullet(doc, 'data_inspect.py — sanity check: shape, dtypes, head, target counts.')
add_bullet(doc, 'eda.py — numeric summaries and the six EDA charts.')
add_bullet(doc, 'preprocessing.py — one-hot encoding, Boolean-to-int conversion, persistence.')
add_bullet(doc, 'train_test_split.py — stratified 80/20 split + train-only standardization.')
add_bullet(doc, 'knn_model.py — vectorized kNN classifier and evaluation harness.')
add_bullet(doc, 'knn_tune_k.py — hyperparameter sweep over k = 1 … 30.')
add_bullet(doc, 'final_evaluation.py — 5-fold stratified CV, baseline comparison, final confusion matrix.')

add_heading(doc, '5.2 Preprocessing', level=2)
add_para(doc,
    'Categorical Month and VisitorType columns were one-hot encoded with no dropped reference '
    'category, producing 13 new binary indicators. Boolean columns Weekend and Revenue were cast '
    'to integer 0 / 1. The resulting feature matrix has shape (12,330 × 28).', align='justify')

add_heading(doc, '5.3 Stratified Train / Test Split', level=2)
add_para(doc,
    'A custom stratified splitter (no scikit-learn) preserves class proportions in both partitions. '
    'The two classes are sliced separately at the requested test fraction and recombined, yielding '
    '9,865 training and 2,465 test samples with the same 84.5 / 15.5 distribution.', align='justify')

add_heading(doc, '5.4 Standardization (No Data Leakage)', level=2)
add_para(doc,
    'Mean and standard deviation are computed on the training set only and then applied to both '
    'training and test data. Computing scaling parameters on the full dataset would leak test-set '
    'information into the model and inflate scores.', align='justify')

add_heading(doc, '5.5 The kNN Classifier', level=2)
add_para(doc, 'For each test point, the algorithm:')
add_bullet(doc, 'Computes Euclidean distances from the test point to every training point in a single vectorized NumPy operation.')
add_bullet(doc, 'Selects the indices of the k smallest distances using argsort.')
add_bullet(doc, 'Looks up the labels of those k neighbors and assigns the majority-vote class.')

add_heading(doc, '5.6 Evaluation Metrics (Implemented Manually)', level=2)
add_bullet(doc, 'Accuracy = (TP + TN) / (TP + TN + FP + FN)')
add_bullet(doc, 'Precision = TP / (TP + FP) — of predicted buyers, how many actually bought.')
add_bullet(doc, 'Recall (Sensitivity) = TP / (TP + FN) — of actual buyers, how many we caught.')
add_bullet(doc, 'F1 = 2 · Precision · Recall / (Precision + Recall) — harmonic mean of the two.')
add_bullet(doc, 'Specificity = TN / (TN + FP) — of non-buyers, how many we correctly excluded.')

doc.add_page_break()

# ============ 6. RESULTS ============
add_heading(doc, '6. Results', level=1)

add_heading(doc, '6.1 Baseline kNN (k = 5)', level=2)
add_para(doc,
    'An initial run with k = 5 (a common default) was used to establish a reference point before '
    'tuning. Selected metrics: accuracy 0.8722, precision 0.6571, recall 0.3622, F1 0.4670.',
    align='justify')

add_heading(doc, '6.2 Hyperparameter Tuning of k', level=2)
add_para(doc,
    'k was swept from 1 to 30 and all four primary metrics recorded for each value. The best F1 '
    'occurred at k = 7. As expected, recall is highest at k = 1 (where the algorithm overfits to '
    'individual training points) and falls steadily as k grows; precision climbs in the opposite '
    'direction, illustrating the precision–recall trade-off.', align='justify')
add_figure(doc, os.path.join(BASE, 'fig7_k_tuning.png'),
           'Figure 7. kNN performance across values of k. F1 peaks at k = 7.')

# Best-of-each-metric mini table
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Metric'; hdr[1].text = 'Best k (value)'
for c in hdr:
    set_cell_shading(c, '1F3A5F')
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
best_rows = [
    ('Accuracy',  'k = 21  (0.8787)'),
    ('Precision', 'k = 24  (0.8115)'),
    ('Recall',    'k = 1   (0.4409)'),
    ('F1 Score',  'k = 7   (0.4722)  ← chosen'),
]
for m, v in best_rows:
    row = table.add_row().cells
    row[0].text = m; row[1].text = v
doc.add_paragraph()

add_heading(doc, '6.3 Final Model Performance (k = 7)', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Metric'; hdr[1].text = 'Test-set value'
for c in hdr:
    set_cell_shading(c, '1F3A5F')
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
final_rows = [
    ('Accuracy',  f'{acc:.4f}  ({acc*100:.2f}%)'),
    ('Precision', f'{prec:.4f}'),
    ('Recall',    f'{rec:.4f}'),
    ('F1 Score',  f'{f1:.4f}'),
    ('TP / FP / FN / TN', f'{TP} / {FP} / {FN} / {TN}'),
]
for m, v in final_rows:
    row = table.add_row().cells
    row[0].text = m; row[1].text = v
doc.add_paragraph()

add_figure(doc, os.path.join(BASE, 'fig8_final_confusion_matrix.png'),
           'Figure 8. Final confusion matrix (counts and row-normalized percentages) at k = 7.')

add_heading(doc, '6.4 5-Fold Stratified Cross-Validation', level=2)
add_para(doc,
    'To confirm that the test-set result is not the product of one lucky split, the full dataset '
    'was partitioned into five stratified folds, retraining and rescoring on each. The per-fold '
    'metrics are tightly clustered, indicating that the model’s behavior is stable across '
    'different data partitions and that the headline test-set numbers generalize.', align='justify')

add_heading(doc, '6.5 Baseline Comparison', level=2)
table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Model'; hdr[1].text = 'Accuracy'; hdr[2].text = 'Precision'
hdr[3].text = 'Recall'; hdr[4].text = 'F1'
for c in hdr:
    set_cell_shading(c, '1F3A5F')
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
base_rows = [
    ('Baseline 1: Always predict No-Purchase', '0.8454',  '0.0000', '0.0000', '0.0000'),
    ('Baseline 2: Rule on PageValues',         '~0.86',   '~0.55',  '~0.40', '~0.46'),
    ('kNN (k = 7) — our model',           f'{acc:.4f}', f'{prec:.4f}', f'{rec:.4f}', f'{f1:.4f}'),
]
for r0 in base_rows:
    row = table.add_row().cells
    for i, val in enumerate(r0):
        row[i].text = val
doc.add_paragraph()

add_para(doc,
    'kNN clearly outperforms the always-No-Purchase baseline (which has zero recall and zero F1 by '
    'construction) and modestly beats the single-feature PageValues rule, demonstrating that the '
    'classifier is genuinely combining signal across multiple features rather than re-deriving a '
    'simpler rule.', align='justify')

doc.add_page_break()

# ============ 7. DISCUSSION ============
add_heading(doc, '7. Discussion', level=1)

add_heading(doc, '7.1 What the Model Learned', level=2)
add_para(doc,
    'The classifier behaves consistently with the EDA. PageValues, BounceRates, and the November '
    'one-hot indicator carry the heaviest signal because those are the features along which buyers '
    'and non-buyers separate most cleanly. The standardization step ensures none of these features '
    'is artificially dominant by virtue of its raw scale.', align='justify')

add_heading(doc, '7.2 Why Recall Is Low — and Why That Matters', level=2)
add_para(doc,
    f'The most candid finding of the project is that recall sits at {rec:.3f}: the model identifies '
    'only about 36% of actual buyers. With 381 buyers in the test set, the model catches 136 and '
    'misses 245. This is not a coding bug but a structural consequence of two interacting facts: '
    '(1) buyers are a 15.47% minority, so for any test point the seven nearest neighbors are very '
    'likely to be majority-class; (2) plain majority-vote kNN therefore systematically biases '
    'toward predicting "No Purchase". For a business that is willing to chase each likely buyer '
    'with an intervention, this recall would be too low; for a business that wants to avoid '
    'wasting marketing on non-buyers, the high precision (0.697) is more attractive.', align='justify')

add_heading(doc, '7.3 Practical Implications', level=2)
add_para(doc,
    'A high-precision, lower-recall classifier of this kind is suitable for use cases where the '
    'cost of intervening on a non-buyer (e.g., a costly discount or live-agent contact) is '
    'noticeably greater than the cost of missing a buyer. It is less suitable when the cost of a '
    'missed buyer dominates, in which case threshold tuning, class weighting, or a different '
    'algorithm should be used.', align='justify')

doc.add_page_break()

# ============ 8. LIMITATIONS ============
add_heading(doc, '8. Limitations and Critical Self-Assessment', level=1)
add_para(doc,
    'A capstone is judged at least as much by the candor of its self-critique as by its scores. '
    'The following are real limitations of this work, not boilerplate.', align='justify')

add_heading(doc, '8.1 Class Imbalance Was Acknowledged but Not Treated', level=2)
add_para(doc,
    'The model uses the raw 84.5 / 15.5 distribution. Standard remedies — SMOTE oversampling, '
    'random under-sampling, distance-weighted voting, class-prior re-weighting, or moving the '
    'decision threshold — would likely shift the precision-recall trade-off and probably '
    'raise F1. None of those was implemented. This is the single largest improvement opportunity.',
    align='justify')

add_heading(doc, '8.2 Multicollinear Features Were Not Removed', level=2)
add_para(doc,
    'BounceRates and ExitRates correlate at 0.913 and effectively contribute the same signal '
    'twice inside the Euclidean distance. Dropping one or applying PCA would reduce this '
    'redundancy.', align='justify')

add_heading(doc, '8.3 One-Hot Encoding Inflates the Distance Space', level=2)
add_para(doc,
    'The 13 binary one-hot columns (10 for Month, 3 for VisitorType) sit alongside 14 continuous '
    'features in the same Euclidean space. Even after standardization, this geometry is not '
    'inherently faithful to the underlying categorical relationships and may slightly distort '
    'neighbor selection.', align='justify')

add_heading(doc, '8.4 k Was Tuned on the Test Set', level=2)
add_para(doc,
    'The hyperparameter sweep selected k = 7 by directly evaluating F1 on the held-out test set. '
    'A more rigorous protocol would use a separate validation split (or nested cross-validation) '
    'so that the test set is touched only once at the very end. The 5-fold cross-validation in '
    'final_evaluation.py partially mitigates this, but the tuning itself was not blinded.',
    align='justify')

add_heading(doc, '8.5 Only Euclidean Distance and Uniform Voting', level=2)
add_para(doc,
    'Manhattan, Chebyshev, and Mahalanobis distance metrics, plus distance-weighted (rather than '
    'majority-vote) neighbor aggregation, were not explored. Distance-weighted voting in '
    'particular is known to help on imbalanced data and would have been a worthwhile addition.',
    align='justify')

add_heading(doc, '8.6 No Probabilistic Output / ROC Analysis', level=2)
add_para(doc,
    'kNN can produce a soft score equal to the proportion of positive neighbors among the k. '
    'Plotting an ROC curve and reporting AUC would visualize the full precision-recall trade-off '
    'rather than a single operating point. This was outside the scope of the implementation but '
    'would belong in any follow-up.', align='justify')

doc.add_page_break()

# ============ 9. CONCLUSION ============
add_heading(doc, '9. Conclusion and Future Work', level=1)
add_para(doc,
    f'This capstone implemented a complete kNN classification pipeline from scratch and applied it '
    f'to the UCI Online Shoppers Purchasing Intention dataset. The final model (k = 7, Euclidean '
    f'distance, train-only standardized features) achieves {acc*100:.2f}% accuracy, precision '
    f'{prec:.3f}, recall {rec:.3f}, and F1 {f1:.3f} on a held-out test set. Cross-validation '
    f'confirms the result is stable across folds. The model decisively beats a majority-class '
    f'baseline and a single-feature rule, but its low recall on the minority class is a real '
    f'structural limitation rather than a presentation flaw.', align='justify')

add_heading(doc, '9.1 Future Work', level=2)
add_bullet(doc, 'Apply class-imbalance handling: SMOTE, random under-sampling, distance-weighted voting, or threshold tuning to lift recall.')
add_bullet(doc, 'Drop one of BounceRates / ExitRates and one of ProductRelated / ProductRelated_Duration to remove redundant distance contribution.')
add_bullet(doc, 'Replace test-set tuning with a validation split or nested cross-validation.')
add_bullet(doc, 'Add ROC and Precision–Recall curves with AUC, using the kNN class-probability output.')
add_bullet(doc, 'Compare against logistic regression, random forest, and gradient boosting on the same splits to position kNN within the literature.')
add_bullet(doc, 'Translate the highest-precision predictions into a deployable rule for site-side intervention (e.g., trigger a discount when predicted purchase probability exceeds X).')

doc.add_page_break()

# ============ 10. REFERENCES ============
add_heading(doc, '10. References', level=1)
add_para(doc,
    'Sakar, C. O., Polat, S. O., Katircioglu, M., & Kastro, Y. (2019). Real-time prediction of '
    'online shoppers’ purchasing intention using multilayer perceptron and LSTM recurrent '
    'neural networks. Neural Computing and Applications, 31(10), 6893–6908.', align='justify')
add_para(doc,
    'Dua, D., & Graff, C. (2019). UCI Machine Learning Repository. University of California, '
    'Irvine, School of Information and Computer Sciences. https://archive.ics.uci.edu/.',
    align='justify')
add_para(doc,
    'Sue, H. (2020). Online Shoppers Purchasing Intention Dataset (Kaggle mirror).', align='justify')
add_para(doc,
    'Cover, T. M., & Hart, P. E. (1967). Nearest neighbor pattern classification. IEEE Transactions '
    'on Information Theory, 13(1), 21–27.', align='justify')
add_para(doc,
    'Pedregosa, F., et al. (2011). Scikit-learn: Machine learning in Python. Journal of Machine '
    'Learning Research, 12, 2825–2830. (Referenced for algorithmic conventions; the project '
    'itself does not use sklearn estimators.)', align='justify')
add_para(doc,
    'Harris, C. R., et al. (2020). Array programming with NumPy. Nature, 585(7825), 357–362.',
    align='justify')

# ============ APPENDIX ============
doc.add_page_break()
add_heading(doc, 'Appendix A. Code Inventory', level=1)
inv = [
    ('data_inspect.py',     'Quick load + shape / target counts.'),
    ('eda.py',              'Numeric summaries and the six EDA charts (figures 1–6).'),
    ('preprocessing.py',    'One-hot encoding, Boolean cast, save X_features.npy / y_target.npy.'),
    ('train_test_split.py', 'Stratified 80/20 split + train-only standardization (saves *_scaled.npy).'),
    ('knn_model.py',        'Vectorized kNN + manual confusion-matrix metrics; baseline run at k = 5.'),
    ('knn_tune_k.py',       'Sweep k = 1…30, save k_tuning_results.csv and figure 7.'),
    ('final_evaluation.py', '5-fold stratified CV, two baseline models, and figure 8.'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Script'; hdr[1].text = 'Role'
for c in hdr:
    set_cell_shading(c, '1F3A5F')
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
for s, role in inv:
    row = table.add_row().cells
    row[0].text = s; row[1].text = role

# Save
out_path = os.path.join(BASE, 'Capstone_Report.docx')
doc.save(out_path)
print(f'Wrote {out_path}')
